from __future__ import annotations

import base64
import io
import time

import pytest
from docx import Document as DocxDocument
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.exceptions.template import TemplateRenderError
from app.models.sotto_cartella import SottoCartella
from app.models.template import Template
from app.services.render.docx_walker import build_docx
from app.services.render.html_renderer import build_html
from app.services.render.images import ImageAsset
from app.services.render.pdf_renderer import build_pdf
from app.services.template_service import (
    _default_filename,
    _resolve_filename,
    _sanitize_filename,
)

# PNG 1x1 pixel valido, usato come contenuto immagine minimale nei test.
_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY"
    "42YAAAAASUVORK5CYII="
)


def _read_paragraphs(content: bytes) -> list[str]:
    doc = DocxDocument(io.BytesIO(content))
    return [p.text for p in doc.paragraphs]


def test_build_docx_paragraph_and_heading():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Titolo"}],
            },
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            },
        ],
    }
    content = build_docx(contenuto, {})
    assert _read_paragraphs(content) == ["Titolo", "Testo semplice"]


def test_build_docx_bold_and_italic_marks():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": "grassetto",
                        "marks": [{"type": "bold"}],
                    },
                    {"type": "text", "text": " e "},
                    {
                        "type": "text",
                        "text": "corsivo",
                        "marks": [{"type": "italic"}],
                    },
                ],
            }
        ],
    }
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    runs = doc.paragraphs[0].runs
    assert runs[0].text == "grassetto"
    assert runs[0].bold is True
    assert runs[1].text == " e "
    assert not runs[1].bold
    assert runs[2].text == "corsivo"
    assert runs[2].italic is True


def test_build_docx_mergefield_resolved():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Gentile "},
                    {
                        "type": "mergefield",
                        "attrs": {"chiave": "socio.nome"},
                    },
                ],
            }
        ],
    }
    content = build_docx(contenuto, {"socio": {"nome": "Mario"}})
    assert _read_paragraphs(content) == ["Gentile Mario"]


def test_build_docx_mergefield_non_risolvibile():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {
                        "type": "mergefield",
                        "attrs": {"chiave": "socio.nome"},
                    }
                ],
            }
        ],
    }
    content = build_docx(contenuto, {})
    assert _read_paragraphs(content) == [""]


def test_build_docx_mergefield_entita_assente_dal_context():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Gentile "},
                    {
                        "type": "mergefield",
                        "attrs": {"chiave": "socio.nome"},
                    },
                ],
            }
        ],
    }
    content = build_docx(contenuto, {"altra_entita": {"campo": "x"}})
    assert _read_paragraphs(content) == ["Gentile "]


def test_build_docx_tipo_nodo_non_supportato():
    contenuto = {"type": "doc", "content": [{"type": "codeBlock", "content": []}]}
    with pytest.raises(TemplateRenderError):
        build_docx(contenuto, {})


def test_build_html_risolve_mergefield_e_marks():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Titolo"}],
            },
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Gentile "},
                    {"type": "mergefield", "attrs": {"chiave": "socio.nome"}},
                    {
                        "type": "text",
                        "text": " grassetto",
                        "marks": [{"type": "bold"}],
                    },
                ],
            },
        ],
    }
    html = build_html(contenuto, {"socio": {"nome": "Mario"}})
    assert "<h1>Titolo</h1>" in html
    assert "Gentile Mario" in html
    assert "<strong> grassetto</strong>" in html
    assert "@page { size: A4; margin: 2cm; }" in html
    assert "unpkg.com" not in html
    assert "http" not in html


def test_build_html_mergefield_non_risolvibile():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "mergefield", "attrs": {"chiave": "socio.nome"}}],
            }
        ],
    }
    html = build_html(contenuto, {})
    assert "<p></p>" in html
    assert "campo mancante" not in html


def test_build_html_mergefield_entita_assente_dal_context():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Gentile "},
                    {"type": "mergefield", "attrs": {"chiave": "socio.nome"}},
                ],
            }
        ],
    }
    html = build_html(contenuto, {"altra_entita": {"campo": "x"}})
    assert "<p>Gentile </p>" in html
    assert "campo mancante" not in html


def test_build_html_tipo_nodo_non_supportato():
    contenuto = {"type": "doc", "content": [{"type": "codeBlock", "content": []}]}
    with pytest.raises(TemplateRenderError):
        build_html(contenuto, {})


def _paragrafo_con_align(text_align: str | None) -> dict:
    attrs = {} if text_align is None else {"textAlign": text_align}
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "attrs": attrs,
                "content": [{"type": "text", "text": "Testo"}],
            }
        ],
    }


@pytest.mark.parametrize(
    "text_align,expected",
    [("center", 1), ("right", 2), ("justify", 3)],
)
def test_build_docx_text_align(text_align, expected):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    content = build_docx(_paragrafo_con_align(text_align), {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH(expected)


@pytest.mark.parametrize("text_align", ["center", "right", "justify"])
def test_build_html_text_align(text_align):
    html = build_html(_paragrafo_con_align(text_align), {})
    assert f'style="text-align: {text_align};"' in html


def test_build_docx_heading_text_align():
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1, "textAlign": "center"},
                "content": [{"type": "text", "text": "Titolo"}],
            }
        ],
    }
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_build_html_heading_text_align():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1, "textAlign": "right"},
                "content": [{"type": "text", "text": "Titolo"}],
            }
        ],
    }
    html = build_html(contenuto, {})
    assert '<h1 style="text-align: right;">Titolo</h1>' in html


def _paragrafo_con_text_style(attrs: dict) -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": "Colorato",
                        "marks": [{"type": "textStyle", "attrs": attrs}],
                    }
                ],
            }
        ],
    }


def test_build_docx_text_style_color_e_font():
    from docx.shared import RGBColor

    contenuto = _paragrafo_con_text_style({"color": "#FF0000", "fontFamily": "Arial"})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor.from_string("FF0000")
    assert run.font.name == "Arial"


def test_build_html_text_style_color_e_font():
    contenuto = _paragrafo_con_text_style({"color": "#FF0000", "fontFamily": "Arial"})
    html = build_html(contenuto, {})
    assert '<span style="color: #FF0000; font-family: Arial;">Colorato</span>' in html


def test_build_html_text_style_solo_color():
    contenuto = _paragrafo_con_text_style({"color": "#00FF00"})
    html = build_html(contenuto, {})
    assert '<span style="color: #00FF00;">Colorato</span>' in html


def test_build_docx_text_style_font_non_whitelisted_ignorato():
    contenuto = _paragrafo_con_text_style({"fontFamily": "ComicSansMS"})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.name is None


def test_build_html_text_style_font_non_whitelisted_ignorato():
    contenuto = _paragrafo_con_text_style({"fontFamily": "ComicSansMS"})
    html = build_html(contenuto, {})
    assert "ComicSansMS" not in html
    assert "<span" not in html


def test_build_docx_senza_textalign_ne_textstyle_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Titolo"}],
            },
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "grassetto", "marks": [{"type": "bold"}]}
                ],
            },
        ],
    }
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].alignment is None
    assert doc.paragraphs[1].alignment is None
    assert doc.paragraphs[1].runs[0].bold is True


def test_build_html_senza_textalign_ne_textstyle_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Titolo"}],
            },
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "grassetto", "marks": [{"type": "bold"}]}
                ],
            },
        ],
    }
    html = build_html(contenuto, {})
    assert "<h1>Titolo</h1>" in html
    assert "<p><strong>grassetto</strong></p>" in html
    assert "style=" not in html


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_docx_text_style_colore_malformato_ignorato(colore_malformato):
    contenuto = _paragrafo_con_text_style({"color": colore_malformato})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.color.rgb is None


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_html_text_style_colore_malformato_ignorato(colore_malformato):
    contenuto = _paragrafo_con_text_style({"color": colore_malformato})
    html = build_html(contenuto, {})
    assert "color:" not in html
    assert "<span" not in html or "font-family:" in html


def test_build_docx_text_style_colore_valido_con_e_senza_hash():
    for colore in ["#FF0000", "FF0000"]:
        contenuto = _paragrafo_con_text_style({"color": colore})
        content = build_docx(contenuto, {})
        doc = DocxDocument(io.BytesIO(content))
        run = doc.paragraphs[0].runs[0]
        from docx.shared import RGBColor

        assert run.font.color.rgb == RGBColor.from_string("FF0000")


def test_build_html_text_style_colore_valido_con_e_senza_hash():
    for colore in ["#FF0000", "FF0000"]:
        contenuto = _paragrafo_con_text_style({"color": colore})
        html = build_html(contenuto, {})
        assert "color: #FF0000" in html


def test_build_docx_text_style_font_size_valido():
    from docx.shared import Pt

    contenuto = _paragrafo_con_text_style({"fontSize": 14})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.size == Pt(14)


def test_build_html_text_style_font_size_valido():
    contenuto = _paragrafo_con_text_style({"fontSize": 14})
    html = build_html(contenuto, {})
    assert '<span style="font-size: 14.0pt;">Colorato</span>' in html


@pytest.mark.parametrize("font_size_fuori_range", [200, 2])
def test_build_docx_text_style_font_size_fuori_range_ignorato(font_size_fuori_range):
    contenuto = _paragrafo_con_text_style({"fontSize": font_size_fuori_range})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.size is None


@pytest.mark.parametrize("font_size_fuori_range", [200, 2])
def test_build_html_text_style_font_size_fuori_range_ignorato(font_size_fuori_range):
    contenuto = _paragrafo_con_text_style({"fontSize": font_size_fuori_range})
    html = build_html(contenuto, {})
    assert "font-size:" not in html
    assert "<span" not in html


@pytest.mark.parametrize("font_size_malformato", ["14pt", None, [], {}])
def test_build_docx_text_style_font_size_malformato_ignorato(font_size_malformato):
    contenuto = _paragrafo_con_text_style({"fontSize": font_size_malformato})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.size is None


@pytest.mark.parametrize("font_size_malformato", ["14pt", None, [], {}])
def test_build_html_text_style_font_size_malformato_ignorato(font_size_malformato):
    contenuto = _paragrafo_con_text_style({"fontSize": font_size_malformato})
    html = build_html(contenuto, {})
    assert "font-size:" not in html
    assert "<span" not in html


def test_build_docx_text_style_color_font_e_font_size_insieme():
    from docx.shared import Pt, RGBColor

    contenuto = _paragrafo_con_text_style(
        {"color": "#FF0000", "fontFamily": "Arial", "fontSize": 18}
    )
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor.from_string("FF0000")
    assert run.font.name == "Arial"
    assert run.font.size == Pt(18)


def test_build_html_text_style_color_font_e_font_size_insieme():
    contenuto = _paragrafo_con_text_style(
        {"color": "#FF0000", "fontFamily": "Arial", "fontSize": 18}
    )
    html = build_html(contenuto, {})
    expected = (
        '<span style="color: #FF0000; font-family: Arial; '
        'font-size: 18.0pt;">Colorato</span>'
    )
    assert expected in html


def test_build_docx_text_style_senza_font_size_nessuna_regressione():
    contenuto = _paragrafo_con_text_style({"color": "#FF0000"})
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    run = doc.paragraphs[0].runs[0]
    assert run.font.size is None


def test_build_html_text_style_senza_font_size_nessuna_regressione():
    contenuto = _paragrafo_con_text_style({})
    html = build_html(contenuto, {})
    assert "<span" not in html
    assert html.count("Colorato") == 1


def _lista_semplice(list_type: str, extra_attrs: dict | None = None) -> dict:
    attrs = extra_attrs or {}
    return {
        "type": "doc",
        "content": [
            {
                "type": list_type,
                "attrs": attrs,
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "Primo"}],
                            }
                        ],
                    },
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "Secondo"}],
                            }
                        ],
                    },
                ],
            }
        ],
    }


def _lista_annidata_due_livelli(list_type: str) -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": list_type,
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "Livello 0"}],
                            },
                            {
                                "type": list_type,
                                "content": [
                                    {
                                        "type": "listItem",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [
                                                    {
                                                        "type": "text",
                                                        "text": "Livello 1",
                                                    }
                                                ],
                                            }
                                        ],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            }
        ],
    }


def test_build_docx_lista_puntata_semplice():
    content = build_docx(_lista_semplice("bulletList"), {})
    doc = DocxDocument(io.BytesIO(content))
    assert [p.text for p in doc.paragraphs] == ["Primo", "Secondo"]
    assert doc.paragraphs[0].style.name == "List Bullet"
    assert doc.paragraphs[1].style.name == "List Bullet"


def test_build_html_lista_puntata_semplice():
    html = build_html(_lista_semplice("bulletList"), {})
    assert '<ul style="list-style-type: disc;">' in html
    assert "<li><p>Primo</p></li>" in html
    assert "<li><p>Secondo</p></li>" in html


def test_build_docx_lista_puntata_annidata_marcatori_per_livello():
    content = build_docx(_lista_annidata_due_livelli("bulletList"), {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].text == "Livello 0"
    assert doc.paragraphs[0].style.name == "List Bullet"
    assert doc.paragraphs[1].text == "Livello 1"
    assert doc.paragraphs[1].style.name == "List Bullet 2"


def test_build_html_lista_puntata_annidata_marcatori_per_livello():
    html = build_html(_lista_annidata_due_livelli("bulletList"), {})
    assert '<ul style="list-style-type: disc;">' in html
    assert '<ul style="list-style-type: circle;">' in html


def test_build_docx_lista_numerata_semplice():
    content = build_docx(_lista_semplice("orderedList"), {})
    doc = DocxDocument(io.BytesIO(content))
    assert [p.text for p in doc.paragraphs] == ["Primo", "Secondo"]
    assert doc.paragraphs[0].style.name == "List Number"
    assert doc.paragraphs[1].style.name == "List Number"


def test_build_html_lista_numerata_semplice():
    html = build_html(_lista_semplice("orderedList"), {})
    assert '<ol style="list-style-type: decimal;">' in html


def test_build_docx_lista_numerata_annidata_marcatori_per_livello():
    content = build_docx(_lista_annidata_due_livelli("orderedList"), {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].style.name == "List Number"
    assert doc.paragraphs[1].style.name == "List Number 2"


def test_build_html_lista_numerata_annidata_marcatori_per_livello():
    html = build_html(_lista_annidata_due_livelli("orderedList"), {})
    assert '<ol style="list-style-type: decimal;">' in html
    assert '<ol style="list-style-type: lower-alpha;">' in html


def test_build_docx_lista_numerata_start_personalizzato():
    content = build_docx(_lista_semplice("orderedList", {"start": 5}), {})
    doc = DocxDocument(io.BytesIO(content))
    numbering_xml = doc.part.numbering_part.element.xml
    assert 'w:startOverride w:val="5"' in numbering_xml


def test_build_html_lista_numerata_start_personalizzato():
    html = build_html(_lista_semplice("orderedList", {"start": 5}), {})
    assert '<ol start="5"' in html


def test_build_docx_lista_item_con_bold_e_mergefield():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "grassetto",
                                        "marks": [{"type": "bold"}],
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "mergefield",
                                        "attrs": {"chiave": "socio.nome"},
                                    }
                                ],
                            }
                        ],
                    },
                ],
            }
        ],
    }
    content = build_docx(contenuto, {"socio": {"nome": "Mario"}})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].runs[0].bold is True
    assert doc.paragraphs[1].text == "Mario"


def test_build_html_lista_item_con_bold_e_mergefield():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "grassetto",
                                        "marks": [{"type": "bold"}],
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "mergefield",
                                        "attrs": {"chiave": "socio.nome"},
                                    }
                                ],
                            }
                        ],
                    },
                ],
            }
        ],
    }
    html = build_html(contenuto, {"socio": {"nome": "Mario"}})
    assert "<li><p><strong>grassetto</strong></p></li>" in html
    assert "<li><p>Mario</p></li>" in html


def test_build_docx_lista_senza_liste_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    content = build_docx(contenuto, {})
    assert _read_paragraphs(content) == ["Testo semplice"]


def test_build_html_lista_senza_liste_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    html = build_html(contenuto, {})
    assert "<ul" not in html
    assert "<ol" not in html
    assert "<p>Testo semplice</p>" in html


def _tableCell(testo: str, *, header: bool = False, attrs: dict | None = None) -> dict:
    node = {
        "type": "tableHeader" if header else "tableCell",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": testo}]}
        ],
    }
    if attrs:
        node["attrs"] = attrs
    return node


def _tabella_semplice_2x2() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            _tableCell("Nome", header=True),
                            _tableCell("Età", header=True),
                        ],
                    },
                    {
                        "type": "tableRow",
                        "content": [
                            _tableCell("Mario"),
                            _tableCell("30"),
                        ],
                    },
                ],
            }
        ],
    }


def test_build_docx_tabella_semplice_2x2():
    content = build_docx(_tabella_semplice_2x2(), {})
    doc = DocxDocument(io.BytesIO(content))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert [c.text for c in table.rows[0].cells] == ["Nome", "Età"]
    assert [c.text for c in table.rows[1].cells] == ["Mario", "30"]


def test_build_html_tabella_semplice_2x2():
    html = build_html(_tabella_semplice_2x2(), {})
    assert "<table" in html
    assert "<th" in html and "<p>Nome</p></th>" in html
    assert "<p>Età</p></th>" in html
    assert "<p>Mario</p></td>" in html
    assert "<p>30</p></td>" in html


def _tabella_con_row_height() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "attrs": {"height": 60},
                        "content": [
                            _tableCell("Nome", header=True),
                            _tableCell("Età", header=True),
                        ],
                    },
                    {
                        "type": "tableRow",
                        "content": [
                            _tableCell("Mario"),
                            _tableCell("30"),
                        ],
                    },
                ],
            }
        ],
    }


def test_build_docx_tabella_con_row_height():
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import Emu

    content = build_docx(_tabella_con_row_height(), {})
    doc = DocxDocument(io.BytesIO(content))
    table = doc.tables[0]
    assert table.rows[0].height == Emu(60 * 9525)
    assert table.rows[0].height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST


def test_build_docx_tabella_senza_row_height_nessuna_regressione():
    content = build_docx(_tabella_semplice_2x2(), {})
    doc = DocxDocument(io.BytesIO(content))
    table = doc.tables[0]
    assert table.rows[0].height is None
    assert table.rows[0].height_rule is None


def test_build_html_tabella_con_row_height():
    html = build_html(_tabella_con_row_height(), {})
    assert '<tr style="height: 60px;">' in html


def test_build_html_tabella_senza_row_height_nessuna_regressione():
    html = build_html(_tabella_semplice_2x2(), {})
    assert "<tr>" in html
    assert "height:" not in html


def _tabella_con_colspan_rowspan() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            _tableCell("P", attrs={"rowspan": 2}),
                            _tableCell("Q", attrs={"colspan": 2}),
                        ],
                    },
                    {
                        "type": "tableRow",
                        "content": [
                            _tableCell("R"),
                            _tableCell("S"),
                        ],
                    },
                ],
            }
        ],
    }


def test_build_docx_tabella_con_colspan_rowspan():
    content = build_docx(_tabella_con_colspan_rowspan(), {})
    doc = DocxDocument(io.BytesIO(content))
    table = doc.tables[0]
    assert len(table.columns) == 3
    assert table.cell(0, 0).text == "P"
    assert table.cell(1, 0).text == "P"  # cella coperta dal rowspan
    assert table.cell(0, 0)._tc is table.cell(1, 0)._tc  # stessa cella unita
    assert table.cell(0, 1).text == "Q"
    assert table.cell(0, 2).text == "Q"
    assert table.cell(0, 1)._tc is table.cell(0, 2)._tc
    assert table.cell(1, 1).text == "R"
    assert table.cell(1, 2).text == "S"


def test_build_html_tabella_con_colspan_rowspan():
    html = build_html(_tabella_con_colspan_rowspan(), {})
    assert 'rowspan="2"' in html
    assert 'colspan="2"' in html
    assert "<p>P</p>" in html
    assert "<p>Q</p>" in html
    assert "<p>R</p>" in html
    assert "<p>S</p>" in html


def _tabella_con_testo_formattato_e_mergefield() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {
                                "type": "tableCell",
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": "Nome: ",
                                                "marks": [{"type": "bold"}],
                                            },
                                            {
                                                "type": "mergefield",
                                                "attrs": {"chiave": "socio.nome"},
                                            },
                                        ],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ],
    }


def test_build_docx_tabella_cella_con_bold_e_mergefield():
    content = build_docx(
        _tabella_con_testo_formattato_e_mergefield(), {"socio": {"nome": "Mario"}}
    )
    doc = DocxDocument(io.BytesIO(content))
    cell = doc.tables[0].cell(0, 0)
    runs = cell.paragraphs[0].runs
    assert runs[0].text == "Nome: "
    assert runs[0].bold is True
    assert runs[1].text == "Mario"
    assert not runs[1].bold


def test_build_html_tabella_cella_con_bold_e_mergefield():
    html = build_html(
        _tabella_con_testo_formattato_e_mergefield(), {"socio": {"nome": "Mario"}}
    )
    assert "<strong>Nome: </strong>Mario" in html


def _tabella_1x1() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [_tableCell("Riquadro")],
                    }
                ],
            }
        ],
    }


def test_build_docx_tabella_1x1_ha_bordo():
    content = build_docx(_tabella_1x1(), {})
    doc = DocxDocument(io.BytesIO(content))
    table = doc.tables[0]
    assert table.style.name == "Table Grid"
    assert table.cell(0, 0).text == "Riquadro"


def test_build_html_tabella_1x1_ha_bordo():
    html = build_html(_tabella_1x1(), {})
    assert "border-collapse: collapse;" in html
    assert "border: 1px solid #000;" in html


def _tabella_con_cella_stile(attrs: dict) -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [_tableCell("Cella", attrs=attrs)],
                    }
                ],
            }
        ],
    }


@pytest.mark.parametrize(
    "align,expected", [("center", 1), ("right", 2), ("justify", 3)]
)
def test_build_docx_tabella_cella_align(align, expected):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    content = build_docx(_tabella_con_cella_stile({"align": align}), {})
    doc = DocxDocument(io.BytesIO(content))
    cell = doc.tables[0].cell(0, 0)
    assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH(expected)


@pytest.mark.parametrize("align", ["center", "right", "justify"])
def test_build_html_tabella_cella_align(align):
    html = build_html(_tabella_con_cella_stile({"align": align}), {})
    assert f"text-align: {align};" in html


def test_build_docx_tabella_cella_bordo_personalizzato():
    content = build_docx(
        _tabella_con_cella_stile({"borderColor": "#FF0000", "borderWidth": 3}), {}
    )
    doc = DocxDocument(io.BytesIO(content))
    xml = doc.tables[0].cell(0, 0)._tc.xml
    assert 'w:color="FF0000"' in xml
    assert 'w:sz="18"' in xml  # 3px * 6 ottavi di punto per px


def test_build_html_tabella_cella_bordo_personalizzato():
    html = build_html(
        _tabella_con_cella_stile({"borderColor": "#FF0000", "borderWidth": 3}), {}
    )
    assert "border: 3px solid #FF0000;" in html


def test_build_docx_tabella_cella_sfondo():
    content = build_docx(_tabella_con_cella_stile({"backgroundColor": "#00FF00"}), {})
    doc = DocxDocument(io.BytesIO(content))
    xml = doc.tables[0].cell(0, 0)._tc.xml
    assert 'w:fill="00FF00"' in xml


def test_build_html_tabella_cella_sfondo():
    html = build_html(_tabella_con_cella_stile({"backgroundColor": "#00FF00"}), {})
    assert "background-color: #00FF00;" in html


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_docx_tabella_cella_bordo_colore_malformato_ignorato(colore_malformato):
    content = build_docx(
        _tabella_con_cella_stile({"borderColor": colore_malformato}), {}
    )
    doc = DocxDocument(io.BytesIO(content))
    xml = doc.tables[0].cell(0, 0)._tc.xml
    assert "tcBorders" not in xml


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_html_tabella_cella_bordo_colore_malformato_ignorato(colore_malformato):
    html = build_html(_tabella_con_cella_stile({"borderColor": colore_malformato}), {})
    assert "border: 1px solid #000;" in html


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_docx_tabella_cella_sfondo_malformato_ignorato(colore_malformato):
    content = build_docx(
        _tabella_con_cella_stile({"backgroundColor": colore_malformato}), {}
    )
    doc = DocxDocument(io.BytesIO(content))
    xml = doc.tables[0].cell(0, 0)._tc.xml
    assert "shd" not in xml


@pytest.mark.parametrize("colore_malformato", ["notacolor", "#fff", "12345", ""])
def test_build_html_tabella_cella_sfondo_malformato_ignorato(colore_malformato):
    html = build_html(
        _tabella_con_cella_stile({"backgroundColor": colore_malformato}), {}
    )
    assert "background-color:" not in html


def test_build_docx_tabella_cella_senza_stile_nessuna_regressione():
    content = build_docx(_tabella_semplice_2x2(), {})
    doc = DocxDocument(io.BytesIO(content))
    xml = doc.tables[0].cell(0, 0)._tc.xml
    assert "tcBorders" not in xml
    assert "shd" not in xml


def test_build_html_tabella_cella_senza_stile_nessuna_regressione():
    html = build_html(_tabella_semplice_2x2(), {})
    assert "border: 1px solid #000;" in html
    assert "background-color:" not in html
    assert "text-align:" not in html


def test_build_docx_senza_tabelle_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    content = build_docx(contenuto, {})
    doc = DocxDocument(io.BytesIO(content))
    assert doc.tables == []
    assert _read_paragraphs(content) == ["Testo semplice"]


def test_build_html_senza_tabelle_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    html = build_html(contenuto, {})
    assert "<table" not in html
    assert "<p>Testo semplice</p>" in html


def _contenuto_con_immagine(attrs: dict | None = None) -> dict:
    return {
        "type": "doc",
        "content": [
            {"type": "image", "attrs": attrs or {"documentoId": 1}},
        ],
    }


def test_build_docx_immagine_presente():
    images = {1: ImageAsset(content=_PNG_1X1, mime_type="image/png")}
    content = build_docx(_contenuto_con_immagine(), {}, images)
    doc = DocxDocument(io.BytesIO(content))
    assert len(doc.inline_shapes) == 1


def test_build_html_immagine_presente():
    images = {1: ImageAsset(content=_PNG_1X1, mime_type="image/png")}
    html = build_html(_contenuto_con_immagine(), {}, images)
    assert "data:image/png;base64," in html
    b64 = html.split("data:image/png;base64,")[1].split('"')[0]
    assert base64.b64decode(b64) == _PNG_1X1


def test_build_docx_immagine_con_width_e_align():
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu

    images = {1: ImageAsset(content=_PNG_1X1, mime_type="image/png")}
    contenuto = _contenuto_con_immagine(
        {"documentoId": 1, "width": 100, "align": "center"}
    )
    content = build_docx(contenuto, {}, images)
    doc = DocxDocument(io.BytesIO(content))
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    picture = doc.inline_shapes[0]
    assert picture.width == Emu(100 * 9525)


def test_build_html_immagine_con_width_e_align():
    images = {1: ImageAsset(content=_PNG_1X1, mime_type="image/png")}
    contenuto = _contenuto_con_immagine(
        {"documentoId": 1, "width": 100, "align": "center"}
    )
    html = build_html(contenuto, {}, images)
    assert '<p style="text-align: center;">' in html
    assert 'width="100"' in html


def test_build_docx_immagine_non_disponibile_placeholder():
    content = build_docx(_contenuto_con_immagine(), {}, {})
    assert _read_paragraphs(content) == ["[immagine non disponibile]"]


def test_build_html_immagine_non_disponibile_placeholder():
    html = build_html(_contenuto_con_immagine(), {}, {})
    assert "<p>[immagine non disponibile]</p>" in html
    assert "<img" not in html


def test_build_docx_senza_immagini_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    assert build_docx(contenuto, {}, {}) == build_docx(contenuto, {})


def test_build_html_senza_immagini_nessuna_regressione():
    contenuto = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Testo semplice"}],
            }
        ],
    }
    assert build_html(contenuto, {}, {}) == build_html(contenuto, {})


def _contenuto_multi_pagina() -> dict:
    testo_lungo = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 20
    paragrafi = [
        {"type": "paragraph", "content": [{"type": "text", "text": testo_lungo}]}
        for _ in range(80)
    ]
    return {"type": "doc", "content": paragrafi}


@pytest.mark.asyncio
async def test_build_pdf_produce_pdf_valido():
    html = build_html(_contenuto_multi_pagina(), {})

    inizio = time.monotonic()
    content = await build_pdf(html)
    durata = time.monotonic() - inizio

    assert content[:5] == b"%PDF-"
    # Se l'attesa di impaginazione degrada di nuovo a un timeout fisso
    # (invece di risolversi sul completamento reale di paged.js), questo
    # test tornerebbe a durare ~10s+: un tempo basso è la prova che
    # l'hook PagedConfig.after ha davvero segnalato il completamento.
    assert durata < 5, (
        f"build_pdf ha impiegato {durata:.1f}s: l'attesa sembra basarsi "
        "su un timeout invece che sul completamento reale di paged.js"
    )


async def _create_persona(ac: AsyncClient) -> dict:
    resp = await ac.post(
        "/api/v1/persone/",
        json={
            "banda_codice": 1,
            "nome": "Mario",
            "cognome": "Rossi",
            "codice_fiscale": "RSSMRA80A01H501U",
            "data_nascita": "1980-01-01",
        },
    )
    assert resp.status_code == 201, resp.text
    return resp.json()


async def _create_socio_con_template(client: AsyncClient) -> tuple[int, int]:
    ruolo_resp = await client.post(
        "/api/v1/ruoli-banda/", json={"codice": 10, "descrizione": "Socio Bandista"}
    )
    assert ruolo_resp.status_code == 201, ruolo_resp.text

    persona = await _create_persona(client)
    socio_resp = await client.post(
        "/api/v1/soci/",
        json={
            "persona_id": persona["id"],
            "codice_socio": "S001",
            "ruolo_banda_codice": 10,
        },
    )
    assert socio_resp.status_code == 201, socio_resp.text
    socio_id = socio_resp.json()["id"]

    template_resp = await client.post(
        "/api/v1/templates/",
        json={
            "nome": "Modulo Iscrizione",
            "contenuto_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [
                            {"type": "text", "text": "Gentile "},
                            {
                                "type": "mergefield",
                                "attrs": {"chiave": "socio.nome"},
                            },
                            {"type": "text", "text": ", con la presente..."},
                        ],
                    }
                ],
            },
            "entita_richieste": ["socio"],
        },
    )
    assert template_resp.status_code == 201, template_resp.text
    return template_resp.json()["id"], socio_id


@pytest.mark.asyncio
async def test_generate_docx_endpoint_produce_docx_valido(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    documento = generate_resp.json()
    assert documento["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    download_resp = await client.get(f"/api/v1/documenti/{documento['id']}/download")
    assert download_resp.status_code == 200
    doc = DocxDocument(io.BytesIO(download_resp.content))
    assert doc.paragraphs[0].text == "Gentile Mario, con la presente..."


async def _upload_immagine(client: AsyncClient) -> int:
    resp = await client.post(
        "/api/v1/documenti/",
        files=[("file", ("logo.png", _PNG_1X1, "image/png"))],
    )
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


@pytest.mark.asyncio
async def test_generate_docx_endpoint_con_immagine_referenziata(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)
    documento_id = await _upload_immagine(client)

    update_resp = await client.patch(
        f"/api/v1/templates/{template_id}",
        json={
            "contenuto_json": {
                "type": "doc",
                "content": [{"type": "image", "attrs": {"documentoId": documento_id}}],
            }
        },
    )
    assert update_resp.status_code == 200, update_resp.text

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    documento = generate_resp.json()

    download_resp = await client.get(f"/api/v1/documenti/{documento['id']}/download")
    assert download_resp.status_code == 200
    doc = DocxDocument(io.BytesIO(download_resp.content))
    assert len(doc.inline_shapes) == 1


@pytest.mark.asyncio
async def test_generate_pdf_endpoint_produce_pdf_valido(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/pdf",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    documento = generate_resp.json()
    assert documento["mime_type"] == "application/pdf"

    download_resp = await client.get(f"/api/v1/documenti/{documento['id']}/download")
    assert download_resp.status_code == 200
    assert download_resp.content[:5] == b"%PDF-"


@pytest.mark.asyncio
async def test_preview_endpoint_risolve_mergefield(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)

    resp = await client.post(
        f"/api/v1/templates/{template_id}/preview",
        json={
            "contenuto_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [
                            {"type": "text", "text": "Gentile "},
                            {
                                "type": "mergefield",
                                "attrs": {"chiave": "socio.nome"},
                            },
                        ],
                    }
                ],
            },
            "entities": {"socio": socio_id},
        },
    )
    assert resp.status_code == 200, resp.text
    assert "Gentile Mario" in resp.json()["html"]


@pytest.mark.asyncio
async def test_preview_endpoint_usa_contenuto_del_body_non_quello_salvato(
    client: AsyncClient,
):
    template_id, socio_id = await _create_socio_con_template(client)

    resp = await client.post(
        f"/api/v1/templates/{template_id}/preview",
        json={
            "contenuto_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "Contenuto non salvato"}],
                    }
                ],
            },
            "entities": {"socio": socio_id},
        },
    )
    assert resp.status_code == 200, resp.text
    html = resp.json()["html"]
    assert "Contenuto non salvato" in html
    assert "con la presente" not in html


@pytest.mark.asyncio
async def test_preview_endpoint_404_template_inesistente(client: AsyncClient):
    resp = await client.post(
        "/api/v1/templates/999999/preview",
        json={
            "contenuto_json": {"type": "doc", "content": []},
            "entities": {},
        },
    )
    assert resp.status_code == 404


@pytest.mark.asyncio
async def test_preview_endpoint_forbidden_senza_permesso(client: AsyncClient):
    from app.api.deps import get_current_user
    from app.models.utente import TipoUtente, Utente
    from main import app

    def _user_senza_permessi() -> Utente:
        return Utente(id=1, tipo=TipoUtente.UMANO, email="test@example.com")

    app.dependency_overrides[get_current_user] = _user_senza_permessi
    response = await client.post(
        "/api/v1/templates/1/preview",
        json={"contenuto_json": {"type": "doc", "content": []}, "entities": {}},
    )
    assert response.status_code == 403


@pytest.mark.asyncio
async def test_generate_docx_endpoint_forbidden_senza_permesso(client: AsyncClient):
    from app.api.deps import get_current_user
    from app.models.utente import TipoUtente, Utente
    from main import app

    def _user_senza_permessi() -> Utente:
        return Utente(id=1, tipo=TipoUtente.UMANO, email="test@example.com")

    app.dependency_overrides[get_current_user] = _user_senza_permessi
    response = await client.post(
        "/api/v1/templates/1/generate/docx", json={"entities": {}}
    )
    assert response.status_code == 403


@pytest.mark.asyncio
async def test_generate_pdf_endpoint_forbidden_senza_permesso(client: AsyncClient):
    from app.api.deps import get_current_user
    from app.models.utente import TipoUtente, Utente
    from main import app

    def _user_senza_permessi() -> Utente:
        return Utente(id=1, tipo=TipoUtente.UMANO, email="test@example.com")

    app.dependency_overrides[get_current_user] = _user_senza_permessi
    response = await client.post(
        "/api/v1/templates/1/generate/pdf", json={"entities": {}}
    )
    assert response.status_code == 403


# ---------------------------------------------------------------------------
# Nome file (default e personalizzato) + sotto_cartella_id
# ---------------------------------------------------------------------------


def _template(nome: str = "Modulo") -> Template:
    return Template(nome=nome, contenuto_json={}, entita_richieste=[])


def test_sanitize_filename_sostituisce_caratteri_non_validi():
    assert _sanitize_filename('a/b\\c:d*e?f"g<h>i|j') == "a_b_c_d_e_f_g_h_i_j"


def test_sanitize_filename_senza_caratteri_non_validi_nessuna_regressione():
    assert _sanitize_filename("Modulo Iscrizione 2026") == "Modulo Iscrizione 2026"


def test_default_filename_con_socio_nel_context():
    nome = _default_filename(
        _template("Modulo"), {"socio": {"cognome": "Rossi", "nome": "Mario"}}
    )
    assert nome.startswith("Modulo_Rossi_Mario_")
    timestamp = nome.removeprefix("Modulo_Rossi_Mario_")
    assert len(timestamp) == len("AAAAMMGG_HHmm")
    assert timestamp[8] == "_"


def test_default_filename_con_esterno_nel_context_se_socio_assente():
    nome = _default_filename(
        _template("Modulo"), {"esterno": {"cognome": "Bianchi", "nome": "Luca"}}
    )
    assert nome.startswith("Modulo_Bianchi_Luca_")


def test_default_filename_priorita_socio_su_esterno():
    context = {
        "socio": {"cognome": "Rossi", "nome": "Mario"},
        "esterno": {"cognome": "Bianchi", "nome": "Luca"},
    }
    nome = _default_filename(_template("Modulo"), context)
    assert nome.startswith("Modulo_Rossi_Mario_")


def test_default_filename_senza_socio_ne_esterno():
    nome = _default_filename(_template("Modulo"), {})
    assert nome.startswith("Modulo_")
    assert "Modulo__" not in nome


def test_default_filename_sanifica_nome_template_con_caratteri_non_validi():
    nome = _default_filename(_template("Modulo/Iscrizione"), {})
    assert "/" not in nome
    assert nome.startswith("Modulo_Iscrizione_")


def test_resolve_filename_usa_nome_file_esplicito_con_estensione_gia_presente():
    assert _resolve_filename(_template(), {}, "verbale.docx", ".docx") == "verbale.docx"


def test_resolve_filename_appende_estensione_se_assente():
    assert _resolve_filename(_template(), {}, "verbale", ".docx") == "verbale.docx"


def test_resolve_filename_sanifica_nome_file_esplicito():
    assert (
        _resolve_filename(_template(), {}, "verbale/2026", ".pdf") == "verbale_2026.pdf"
    )


def test_resolve_filename_usa_default_se_nome_file_assente():
    nome = _resolve_filename(_template("Modulo"), {}, None, ".pdf")
    assert nome.startswith("Modulo_")
    assert nome.endswith(".pdf")


@pytest.fixture
async def seeded_sotto_cartella(
    seeded_macro_sezioni: None, db_session: AsyncSession
) -> int:
    sc = SottoCartella(nome="Cartella Generati", macro_sezione_codice=1)
    db_session.add(sc)
    await db_session.commit()
    await db_session.refresh(sc)
    return sc.id


@pytest.mark.asyncio
async def test_generate_docx_endpoint_con_nome_file_esplicito(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}, "nome_file": "Modulo Personalizzato"},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    assert generate_resp.json()["nome"] == "Modulo Personalizzato.docx"


@pytest.mark.asyncio
async def test_generate_pdf_endpoint_con_nome_file_esplicito(client: AsyncClient):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/pdf",
        json={"entities": {"socio": socio_id}, "nome_file": "Modulo Personalizzato"},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    assert generate_resp.json()["nome"] == "Modulo Personalizzato.pdf"


@pytest.mark.asyncio
async def test_generate_docx_endpoint_con_nome_file_caratteri_non_validi(
    client: AsyncClient,
):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}, "nome_file": "Modulo/2026:Iscrizione"},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    assert generate_resp.json()["nome"] == "Modulo_2026_Iscrizione.docx"


@pytest.mark.asyncio
async def test_generate_docx_endpoint_senza_nome_file_usa_default(
    client: AsyncClient,
):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    nome = generate_resp.json()["nome"]
    assert nome.startswith("Modulo Iscrizione_Rossi_Mario_")
    assert nome.endswith(".docx")


@pytest.mark.asyncio
async def test_generate_docx_endpoint_con_sotto_cartella_del_template(
    client: AsyncClient, seeded_sotto_cartella: int
):
    template_id, socio_id = await _create_socio_con_template(client)
    update_resp = await client.patch(
        f"/api/v1/templates/{template_id}",
        json={"sotto_cartella_id": seeded_sotto_cartella},
    )
    assert update_resp.status_code == 200, update_resp.text

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    assert generate_resp.json()["sotto_cartella_id"] == seeded_sotto_cartella


@pytest.mark.asyncio
async def test_generate_docx_endpoint_senza_sotto_cartella_retrocompatibile(
    client: AsyncClient,
):
    template_id, socio_id = await _create_socio_con_template(client)

    generate_resp = await client.post(
        f"/api/v1/templates/{template_id}/generate/docx",
        json={"entities": {"socio": socio_id}},
    )
    assert generate_resp.status_code == 200, generate_resp.text
    documento = generate_resp.json()
    assert documento["sotto_cartella_id"] is None
    assert documento["nome"].startswith("Modulo Iscrizione_Rossi_Mario_")
