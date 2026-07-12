from __future__ import annotations

import io
from typing import Any

from docx import Document as DocxDocument
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, RGBColor
from docx.text.paragraph import Paragraph

from app.exceptions.template import TemplateRenderError
from app.services.render.fonts import sanitize_font_family, validate_hex_color
from app.services.render.images import ImageAsset

_HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}

# Schema ProseMirror/TipTap supportato (deve restare sincronizzato con
# html_renderer.py):
# - nodi blocco: doc, paragraph, heading (attrs.level 1-3),
#   bulletList/orderedList (contengono listItem; orderedList ha
#   attrs.start opzionale), listItem (contiene blocchi, tipicamente
#   paragraph, oppure a sua volta bulletList/orderedList annidate)
#   - attrs.textAlign su paragraph/heading: left | center | right | justify
#     (default "left" se assente)
# - nodi tabella: table (contiene tableRow, una o più; nessun attrs
#   rilevante — l'estensione TipTap standard non ne emette di suo),
#   tableRow (contiene tableCell/tableHeader, celle normali e celle di
#   intestazione)
#   - attrs.height (numero opzionale, px) su tableRow: se assente,
#     altezza automatica basata sul contenuto (comportamento invariato)
#   tableCell/tableHeader (contengono blocchi, tipicamente
#   paragraph, ma anche bulletList/orderedList annidate — riusa la logica
#   ricorsiva dei blocchi già esistente)
#   - attrs.colspan (default 1) e attrs.rowspan (default 1) su
#     tableCell/tableHeader: per schema standard ProseMirror/TipTap, le
#     celle "coperte" da un rowspan non compaiono nel JSON delle righe
#     successive (vanno ricostruite camminando la griglia)
#   - attrs.colwidth (array opzionale di larghezze in px) su
#     tableCell/tableHeader
#   - attrs.align (opzionale) su tableCell/tableHeader: stessa
#     whitelist/logica di attrs.textAlign su paragraph/heading, applicata
#     a tutti i paragrafi contenuti nella cella (default nessuno =
#     comportamento attuale)
#   - attrs.borderColor (hex "#RRGGBB", opzionale) e attrs.borderWidth
#     (numero px, opzionale) su tableCell/tableHeader: bordo personalizzato
#     della singola cella; se assenti/non validi resta il bordo di default
#     1px nero della tabella (CR#8)
#   - attrs.backgroundColor (hex "#RRGGBB", opzionale) su tableCell/
#     tableHeader: sfondo della cella, nessuno se assente/non valido
#     (comportamento attuale)
#   - image (blocco, non inline): attrs.documentoId (int, riferimento a un
#     Documento esistente), attrs.width (opzionale, px) e attrs.align
#     (opzionale, left|center|right — stessa whitelist/logica di
#     textAlign/align, default nessuno). Il contenuto binario va risolto a
#     monte (vedi TemplateService._load_images) e passato tramite il
#     parametro images: dict[int, ImageAsset] di build_docx (chiave =
#     documentoId); se un id referenziato non è presente in images (es.
#     Documento cancellato dopo l'inserimento nel template), viene inserito
#     un placeholder testuale "[immagine non disponibile]" invece di far
#     fallire la generazione — stesso principio già usato per i mergefield
#     non risolvibili
# - nodi inline: text, mergefield (attrs.chiave)
# - marks su text/mergefield: bold, italic,
#   textStyle (attrs.color "#RRGGBB", attrs.fontFamily — solo se in
#   app.services.render.fonts.SAFE_FONTS, altrimenti ignorato)
# - marcatori di lista "stile Word": il livello di annidamento (0, 1, 2...)
#   seleziona lo stile built-in di python-docx, riciclando ogni 3 livelli
#   (bulletList: List Bullet / List Bullet 2 / List Bullet 3; orderedList:
#   List Number / List Number 2 / List Number 3). Oltre il terzo livello lo
#   stile si ricicla ma l'indentazione viene forzata manualmente per
#   riflettere la vera profondità.
_TEXT_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_BULLET_LIST_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_ORDERED_LIST_STYLES = ["List Number", "List Number 2", "List Number 3"]
_LIST_INDENT_STEP = Inches(0.25)
_TABLE_STYLE = "Table Grid"
# 1px a 96 DPI (convenzione standard Office/browser) = 9525 EMU.
_EMU_PER_PX = 9525
_BORDER_EDGES = ("top", "left", "bottom", "right")
# w:sz sui bordi è espresso in ottavi di punto: 1px a 96 DPI = 0.75pt = 6 ottavi.
_EIGHTHS_PT_PER_PX = 6
_DEFAULT_BORDER_SZ = _EIGHTHS_PT_PER_PX  # equivalente a 1px, come "Table Grid"
_DEFAULT_BORDER_COLOR = "000000"


def build_docx(
    contenuto_json: dict,
    context: dict,
    images: dict[int, ImageAsset] | None = None,
) -> bytes:
    """Cammina l'albero ProseMirror/TipTap del template e produce un .docx."""
    if contenuto_json.get("type") != "doc":
        raise TemplateRenderError(
            "Il contenuto del template deve avere type 'doc' come radice"
        )

    images = images or {}
    doc = DocxDocument()
    for node in contenuto_json.get("content", []):
        _add_block(doc, node, context, images)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_block(
    doc: Any, node: dict, context: dict, images: dict[int, ImageAsset]
) -> None:
    node_type = node.get("type")
    attrs = node.get("attrs", {})
    if node_type == "paragraph":
        paragraph = doc.add_paragraph()
        _apply_text_align(paragraph, attrs.get("textAlign"))
        _add_inline_content(paragraph, node.get("content", []), context)
    elif node_type == "heading":
        level = attrs.get("level", 1)
        paragraph = doc.add_paragraph(style=_HEADING_STYLES.get(level, "Heading 1"))
        _apply_text_align(paragraph, attrs.get("textAlign"))
        _add_inline_content(paragraph, node.get("content", []), context)
    elif node_type in ("bulletList", "orderedList"):
        _add_list(doc, node, context, depth=0, images=images)
    elif node_type == "table":
        _add_table(doc, node, context, images)
    elif node_type == "image":
        _add_image(doc, node, images)
    else:
        raise TemplateRenderError(f"Tipo di nodo non supportato: {node_type!r}")


def _add_image(doc: Any, node: dict, images: dict[int, ImageAsset]) -> None:
    attrs = node.get("attrs", {})
    asset = images.get(attrs.get("documentoId"))
    paragraph = doc.add_paragraph()
    _apply_text_align(paragraph, attrs.get("align"))
    if asset is None:
        paragraph.add_run("[immagine non disponibile]")
        return
    run = paragraph.add_run()
    width = attrs.get("width")
    if isinstance(width, int | float) and width > 0:
        run.add_picture(io.BytesIO(asset.content), width=Emu(int(width * _EMU_PER_PX)))
    else:
        run.add_picture(io.BytesIO(asset.content))


def _add_list(
    doc: Any, node: dict, context: dict, depth: int, images: dict[int, ImageAsset]
) -> None:
    is_ordered = node.get("type") == "orderedList"
    styles = _ORDERED_LIST_STYLES if is_ordered else _BULLET_LIST_STYLES
    style_name = styles[depth % len(styles)]

    numid_override = None
    if is_ordered:
        start = node.get("attrs", {}).get("start", 1)
        numid_override = _create_numbering_start_override(doc, style_name, start)

    for item in node.get("content", []):
        _add_list_item(doc, item, context, depth, style_name, numid_override, images)


def _add_list_item(
    doc: Any,
    item_node: dict,
    context: dict,
    depth: int,
    style_name: str,
    numid_override: int | None,
    images: dict[int, ImageAsset],
) -> None:
    for child in item_node.get("content", []):
        child_type = child.get("type")
        if child_type == "paragraph":
            paragraph = doc.add_paragraph(style=style_name)
            _apply_list_indent(paragraph, depth)
            if numid_override is not None:
                _set_numid(paragraph, numid_override)
            child_attrs = child.get("attrs", {})
            _apply_text_align(paragraph, child_attrs.get("textAlign"))
            _add_inline_content(paragraph, child.get("content", []), context)
        elif child_type in ("bulletList", "orderedList"):
            _add_list(doc, child, context, depth + 1, images)
        else:
            raise TemplateRenderError(
                f"Tipo di nodo non supportato dentro listItem: {child_type!r}"
            )


def _apply_list_indent(paragraph: Paragraph, depth: int) -> None:
    # I primi 3 livelli usano l'indentazione nativa degli stili built-in
    # (List Bullet/List Bullet 2/List Bullet 3 e le loro varianti
    # numeriche). Oltre quello lo stile si ricicla (vedi _add_list), quindi
    # forziamo un'indentazione crescente per distinguere visivamente la
    # vera profondità di annidamento.
    styles_count = 3
    if depth >= styles_count:
        paragraph.paragraph_format.left_indent = _LIST_INDENT_STEP * (depth + 1)


def _layout_table_rows(node: dict) -> tuple[list[list[tuple[int, dict]]], int]:
    """Calcola, per ogni tableRow, la colonna di partenza di ciascuna cella
    JSON e il numero totale di colonne della tabella.

    Nello schema ProseMirror/TipTap le celle "coperte" da un rowspan
    proveniente da una riga precedente non compaiono nel JSON della riga
    corrente, quindi la griglia va ricostruita tenendo traccia degli
    span ancora attivi riga per riga.
    """
    pending: dict[int, int] = {}
    layout: list[list[tuple[int, dict]]] = []
    total_cols = 0
    for row in node.get("content", []):
        if row.get("type") != "tableRow":
            raise TemplateRenderError(
                f"Tipo di nodo non supportato dentro table: {row.get('type')!r}"
            )
        occupied = set(pending)
        new_pending: dict[int, int] = {
            col: remaining - 1 for col, remaining in pending.items() if remaining > 1
        }
        row_layout: list[tuple[int, dict]] = []
        col = 0
        for cell_node in row.get("content", []):
            cell_type = cell_node.get("type")
            if cell_type not in ("tableCell", "tableHeader"):
                raise TemplateRenderError(
                    f"Tipo di nodo non supportato dentro tableRow: {cell_type!r}"
                )
            while col in occupied:
                col += 1
            attrs = cell_node.get("attrs", {})
            colspan = attrs.get("colspan") or 1
            rowspan = attrs.get("rowspan") or 1
            row_layout.append((col, cell_node))
            if rowspan > 1:
                for c in range(col, col + colspan):
                    new_pending[c] = rowspan - 1
            col += colspan
        row_end = col
        while row_end in occupied:
            row_end += 1
        layout.append(row_layout)
        total_cols = max(total_cols, row_end)
        pending = new_pending
    return layout, total_cols


def _add_table(
    doc: Any, node: dict, context: dict, images: dict[int, ImageAsset]
) -> None:
    layout, total_cols = _layout_table_rows(node)
    num_rows = len(layout)
    if num_rows == 0 or total_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=total_cols, style=_TABLE_STYLE)
    row_nodes = node.get("content", [])
    for row_idx, row_layout in enumerate(layout):
        _apply_row_height(table.rows[row_idx], row_nodes[row_idx].get("attrs", {}))
        for col, cell_node in row_layout:
            attrs = cell_node.get("attrs", {})
            colspan = attrs.get("colspan") or 1
            rowspan = attrs.get("rowspan") or 1
            cell = table.cell(row_idx, col)
            if colspan > 1 or rowspan > 1:
                end_cell = table.cell(row_idx + rowspan - 1, col + colspan - 1)
                cell = cell.merge(end_cell)
            _fill_table_cell(cell, cell_node.get("content", []), context, images)
            _apply_cell_width(cell, attrs.get("colwidth"))
            _apply_cell_align(cell, attrs.get("align"))
            _apply_cell_border_and_background(cell, attrs)


def _fill_table_cell(
    cell: Any, content_nodes: list[dict], context: dict, images: dict[int, ImageAsset]
) -> None:
    # Ogni cella nasce con un paragrafo vuoto già presente (python-docx non
    # permette celle senza paragrafi): lo riutilizziamo solo come segnaposto
    # e lo rimuoviamo se il contenuto reale è stato aggiunto con altri
    # add_paragraph, per non lasciare una riga vuota prima del testo.
    placeholder = cell.paragraphs[0]
    for block in content_nodes:
        _add_block(cell, block, context, images)
    # cell.paragraphs ricostruisce un nuovo Paragraph a ogni accesso, quindi
    # "is" confronterebbe sempre wrapper diversi: si confronta l'elemento
    # XML sottostante, che invece resta lo stesso oggetto.
    if (
        len(cell.paragraphs) > 1
        and cell.paragraphs[0]._p is placeholder._p
        and not placeholder.text
    ):
        _delete_paragraph(placeholder)


def _delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    p.getparent().remove(p)


def _apply_cell_width(cell: Any, colwidth: list[int | None] | None) -> None:
    if not colwidth:
        return
    widths = [w for w in colwidth if w]
    if not widths:
        return
    cell.width = Emu(sum(widths) * _EMU_PER_PX)


def _apply_cell_align(cell: Any, align: str | None) -> None:
    # Si applica a tutti i paragrafi della cella (inclusi quelli di eventuali
    # liste annidate): è un attributo di stile a livello di cella, non del
    # singolo blocco contenuto.
    for paragraph in cell.paragraphs:
        _apply_text_align(paragraph, align)


def _apply_cell_border_and_background(cell: Any, attrs: dict) -> None:
    # python-docx non espone un'API di alto livello per bordo/sfondo della
    # singola cella (solo per lo stile dell'intera tabella), quindi si
    # manipola direttamente l'XML di tcPr.
    border_color = validate_hex_color(attrs.get("borderColor"))
    border_width = attrs.get("borderWidth")
    valid_width = (
        border_width
        if isinstance(border_width, int | float) and border_width > 0
        else None
    )
    background_color = validate_hex_color(attrs.get("backgroundColor"))

    tc_pr = None
    if border_color or valid_width:
        sz = (
            round(valid_width * _EIGHTHS_PT_PER_PX)
            if valid_width
            else _DEFAULT_BORDER_SZ
        )
        color = border_color or _DEFAULT_BORDER_COLOR
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in _BORDER_EDGES:
            edge_elem = OxmlElement(f"w:{edge}")
            edge_elem.set(qn("w:val"), "single")
            edge_elem.set(qn("w:sz"), str(sz))
            edge_elem.set(qn("w:color"), color)
            tc_borders.append(edge_elem)
        tc_pr.append(tc_borders)

    if background_color:
        if tc_pr is None:
            tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), background_color)
        tc_pr.append(shd)


def _apply_row_height(row: Any, attrs: dict) -> None:
    height = attrs.get("height")
    if not height:
        return
    # AT_LEAST (non EXACTLY) replica lo stesso comportamento della resa HTML,
    # dove "height" sul <tr> è un minimo e la riga cresce se il contenuto non
    # ci sta: EXACTLY troncherebbe il contenuto in Word senza equivalente in
    # HTML, disallineando visivamente i due renderer.
    row.height = Emu(height * _EMU_PER_PX)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _create_numbering_start_override(doc: Any, style_name: str, start: int) -> int:
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part.element

    style_elem = doc.styles[style_name].element
    num_id_elem = style_elem.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}")
    base_num_id = num_id_elem.get(qn("w:val"))

    abstract_num_id = None
    for num_elem in numbering_elem.findall(qn("w:num")):
        if num_elem.get(qn("w:numId")) == base_num_id:
            abstract_num_id = num_elem.find(qn("w:abstractNumId")).get(qn("w:val"))
            break

    existing_ids = [
        int(num_elem.get(qn("w:numId")))
        for num_elem in numbering_elem.findall(qn("w:num"))
    ]
    new_num_id = max(existing_ids) + 1

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract_ref)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    lvl_override.append(start_override)
    new_num.append(lvl_override)

    numbering_elem.append(new_num)
    return new_num_id


def _set_numid(paragraph: Paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_elem)
    p_pr.append(num_pr)


def _apply_text_align(paragraph: Paragraph, text_align: str | None) -> None:
    # Se assente (o non riconosciuto) lascia l'allineamento invariato (None),
    # così i template preesistenti senza textAlign producono lo stesso .docx
    # di prima — WD_ALIGN_PARAGRAPH.LEFT lo imposterebbe esplicitamente,
    # cambiando l'XML generato anche se visivamente equivalente.
    if text_align in _TEXT_ALIGN and text_align != "left":
        paragraph.alignment = _TEXT_ALIGN[text_align]


def _add_inline_content(paragraph: Paragraph, nodes: list[dict], context: dict) -> None:
    for node in nodes:
        node_type = node.get("type")
        if node_type == "text":
            _add_run(paragraph, node.get("text", ""), node.get("marks", []))
        elif node_type == "mergefield":
            chiave = node["attrs"]["chiave"]
            _add_run(
                paragraph, _resolve_mergefield(chiave, context), node.get("marks", [])
            )
        else:
            raise TemplateRenderError(
                f"Tipo di nodo inline non supportato: {node_type!r}"
            )


def _add_run(paragraph: Paragraph, text: str, marks: list[dict]) -> None:
    run = paragraph.add_run(text)
    for mark in marks:
        mark_type = mark.get("type")
        if mark_type == "bold":
            run.bold = True
        elif mark_type == "italic":
            run.italic = True
        elif mark_type == "textStyle":
            _apply_text_style(run, mark.get("attrs", {}))


def _apply_text_style(run: Any, attrs: dict) -> None:
    hex_color = validate_hex_color(attrs.get("color"))
    if hex_color:
        run.font.color.rgb = RGBColor.from_string(hex_color)

    font_family = sanitize_font_family(attrs.get("fontFamily"))
    if font_family:
        run.font.name = font_family


def _resolve_mergefield(chiave: str, context: dict) -> str:
    entita, _, campo = chiave.partition(".")
    valori = context.get(entita)
    if not valori or campo not in valori or valori[campo] is None:
        return f"[campo mancante: {chiave}]"
    return str(valori[campo])
